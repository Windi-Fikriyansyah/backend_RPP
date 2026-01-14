from fastapi import APIRouter, Depends, HTTPException, Request, Response
from fastapi.responses import StreamingResponse
import json
import re
import traceback
from app.utils.time_utils import get_jakarta_time
import io
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from app.schemas.rpp_schema import RPPRequest, RPPResponse, RPPData
from app.prompts.rpp_prompt import build_rpp_prompt
from app.gemini_client import gemini_client
from app.security import get_current_user_id # Restored
from app.services.ppt_service import PPTService # Restored

router = APIRouter() # Restored

from app.database import get_db
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from app.models.curriculum import Subject, CurriculumGoal

from datetime import datetime, date

def clean_text(text: str) -> str:
    """Clean text to be compatible with FPDF's default latin-1 fonts."""
    if not text:
        return ""
    # Standardize spaces and common unicode characters
    text = str(text)
    text = text.replace('\r', '')
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2019', "'").replace('\u2018', "'")
    text = text.replace('\u201c', '"').replace('\u201d', '"').replace('\u2022', '\x95')
    # Force to latin-1, replacing anything else with '?'
    # This prevents UnicodeEncodeError in fpdf2
    return text.encode('latin-1', 'replace').decode('latin-1')

def clean_markdown_symbols(text: str) -> str:
    """Extra robust function to remove markdown symbols like **, *, etc."""
    if not text: return ""
    # Remove bold/italic markers
    text = text.replace('***', '').replace('**', '').replace('*', '')
    # Remove underline markers if any
    text = text.replace('__', '').replace('_', '')
    return text.strip()

@router.post("/generate", response_model=RPPResponse)
async def generate_rpp(
    request: RPPRequest, 
    curr_req: Request, 
    user_id: int = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    from app.models.payment import Subscription
    from app.models.rpp_data import SavedRPP, GenerationLog
    from sqlalchemy import func

    # 0. Check Subscription & Limits
    # a. Get active subscription
    sub_res = await db.execute(
        select(Subscription).where(
            Subscription.user_id == user_id,
            Subscription.is_active == True,
            Subscription.end_date > get_jakarta_time()
        )
    )
    subscription = sub_res.scalars().first()
    
    plan_type = subscription.plan_type if subscription else "free"
    
    # b. If Free, check usage
    if plan_type == "free":
        # Count RPP Generations in the current month
        today = date.today()
        first_day = today.replace(day=1)
        
        count_stmt = select(func.count(GenerationLog.id)).where(
            GenerationLog.user_id == user_id,
            GenerationLog.created_at >= first_day
        )
        count_res = await db.execute(count_stmt)
        usage_count = count_res.scalar() or 0
        
        if usage_count >= 3:
            raise HTTPException(
                status_code=403, 
                detail="Kuota generate gratis Anda telah habis (3/3) bulan ini. Silakan upgrade ke paket Pro atau Sekolah untuk akses Unlimited setiap bulan."
            )

    # Proceed with generation...
    # Debug Session
    print(f"DEBUG SESSION: {curr_req.session}")
    
    # 0. Fetch CP Content from DB (Smart Logic)
    db_cp_content = None
    try:
        # Cari CP berdasarkan Mapel (Nama), Fase, dan Elemen
        stmt = select(CurriculumGoal.cp_content).join(Subject).where(
            Subject.name == request.mapel,
            CurriculumGoal.phase == request.fase,
            CurriculumGoal.element == request.elemen
        )
        result = await db.execute(stmt)
        cp_found = result.scalar_one_or_none()
        
        if cp_found:
            db_cp_content = cp_found
        else:
            print(f"CP Not Found for: {request.mapel} - {request.fase} - {request.elemen}")
            
    except Exception as e:
        print(f"Error fetching CP: {e}")

    # 1. Build Prompt with CP
    prompt = build_rpp_prompt(request, db_cp_content)
    
    # 2. Call AI
    result_text = await gemini_client.generate_content(prompt)
    
    # 3. Log Generation (Success)
    new_log = GenerationLog(user_id=user_id, plan_type=plan_type)
    db.add(new_log)
    await db.commit()
    
    # 4. Return
    return RPPResponse(
        data=RPPData(
            rpp_markdown=result_text,
            rpp_json={"note": "Parsed JSON feature coming soon"}
        )
    )

from pydantic import BaseModel
class SaveRPPRequest(BaseModel):
    mapel: str
    kelas: str
    topik: str
    content_markdown: str
    input_data: dict # New field for all form data

@router.post("/save")
async def save_rpp(
    req: SaveRPPRequest,
    db: AsyncSession = Depends(get_db),
    user_id: int = Depends(get_current_user_id)
):
    from app.models.rpp_data import SavedRPP
    
    new_rpp = SavedRPP(
        user_id=user_id,
        mapel=req.mapel,
        kelas=req.kelas,
        topik=req.topik,
        content_markdown=req.content_markdown,
        input_data=req.input_data
    )
    db.add(new_rpp)
    await db.commit()
    await db.refresh(new_rpp)
    
    return {"message": "RPP Saved Successfully", "id": new_rpp.id}

class GeneratePPTRequest(BaseModel):
    rpp_content: str
    mapel: str
    topik: str
    template: str = "auto"

class GenerateQuizRequest(BaseModel):
    rpp_content: str
    mapel: str
    topik: str
    jumlah_soal: int
    tingkat_kesulitan: str

class ExportQuizRequest(BaseModel):
    quiz_data: dict
    mapel: str
    topik: str

class ExportRPPRequest(BaseModel):
    content_markdown: str
    mapel: str
    topik: str
    kelas: str = "Semua"

@router.post("/generate-ppt")
async def generate_ppt_route(
    req: GeneratePPTRequest,
    user_id: int = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    # 1. Check if user is Pro/School
    from app.models.payment import Subscription
    from datetime import datetime
    
    sub_res = await db.execute(
        select(Subscription).where(
            Subscription.user_id == user_id,
            Subscription.is_active == True,
            Subscription.end_date > get_jakarta_time()
        )
    )
    sub = sub_res.scalars().first()
    
    if not sub or sub.plan_type not in ["pro", "school", "monthly", "yearly"]:
        raise HTTPException(status_code=403, detail="Fitur Buat PPT hanya tersedia untuk pelanggan Pro atau Sekolah.")

    # 2. Build Prompt for JSON Structure
    # Determine Theme Instruction
    theme_instruction = """
PILIH TEMA:
Pilih salah satu tema yang paling cocok berdasarkan mapel dan topik:
- "Ceria": Cocok untuk SD, warna oranye/kuning.
- "Formal": Cocok untuk SMP/SMA/Umum, warna biru.
- "Alam": Cocok untuk IPA/Geografi, warna hijau.
- "Pastel": Cocok untuk materi bimbingan atau desain, warna pink/soft.
"""
    if req.template and req.template != "auto":
        theme_instruction = f"""
TEMA DIPILIH USER: "{req.template}"
Gunakan gaya desain "{req.template}" untuk seluruh slide.
"""

    prompt = f"""
Berdasarkan Modul Ajar di bawah ini, buatkan struktur presentasi untuk guru dalam format JSON yang ESTETIK.
Fokus pada materi inti. Buatlah menjadi 7-8 slide.

{theme_instruction}

STRUKTUR JSON:
{{
  "judul_materi": "Judul Besar",
  "theme": "NamaTema",
  "slides": [
    {{
      "judul_slide": "Judul Slide",
      "konten": ["Poin 1", "Poin 2"],
      "keyword_visual": "keyword inggris",
      "layout_type": "split" | "big_image" | "highlight"
    }}
  ]
}}

ATURAN LAYOUT:
- "split": Teks di kiri, gambar di kanan. Gunakan untuk materi standar.
- "big_image": Gambar memenuhi slide dengan judul di bawah. Gunakan untuk visual kuat.
- "highlight": Hanya judul besar di tengah dengan background solid. Gunakan untuk kutipan atau poin kunci.

Campur jenis layout agar tidak bosan.
Format harus JSON murni tanpa teks penjelasan lain.

ISI MODUL AJAR:
{req.rpp_content}
"""

    try:
        # 3. Call AI
        print(f"DEBUG: Generating Slide JSON for {req.topik}...")
        response_text = await gemini_client.generate_content(prompt)
        print(f"DEBUG: Raw AI Response: {response_text[:200]}...")
        
        # Clean JSON: Extract only the part between the first { and the last }
        match = re.search(r'(\{.*\}|\[.*\])', response_text, re.DOTALL)
        if not match:
            print(f"DEBUG: No JSON structure found in response: {response_text}")
            raise HTTPException(status_code=500, detail="AI tidak memberikan format data yang benar.")
            
        clean_json = match.group(0)
        
        try:
            data = json.loads(clean_json)
        except Exception as json_err:
            print(f"DEBUG: JSON Parse Error: {json_err}. Content: {clean_json}")
            raise HTTPException(status_code=500, detail="AI memberikan format JSON yang tidak valid.")
        
        if "slides" not in data:
             print(f"DEBUG: Missing 'slides' key in: {data}")
             raise HTTPException(status_code=500, detail="Data slide tidak lengkap.")

        # 4. Generate PPTX File
        print(f"DEBUG: Generating PPTX File for {len(data.get('slides', []))} slides...")
        ppt_file = await PPTService.generate_ppt(data)
        
        # 5. Return as Download
        # Clean filename from potentially unsafe characters
        safe_topik = re.sub(r'[^\w\s-]', '', req.topik).strip().replace(" ", "_")
        filename = f"PPT_{safe_topik}.pptx"
        print(f"DEBUG: PPTX Generated successfully. Sending {filename}")
        
        return StreamingResponse(
            ppt_file,
            media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            headers={
                "Content-Disposition": f"attachment; filename={filename}",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
        
    except Exception as e:
        import traceback
        traceback.print_exc() # Print full stack trace to terminal
        print(f"Error generating PPT: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Gagal generate PPT: {str(e)}")

@router.post("/generate-quiz")
async def generate_quiz(
    req: GenerateQuizRequest,
    user_id: int = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    from app.models.rpp_data import SavedQuiz
    
    # 1. Build Prompt
    prompt = f"""
Berdasarkan Modul Ajar berikut:
{req.rpp_content}

Buatkan {req.jumlah_soal} soal pilihan ganda dengan tingkat kesulitan {req.tingkat_kesulitan}.

Aturan:
1. Gunakan bahasa Indonesia yang baku dan sesuai umur siswa di Fase tersebut.
2. Berikan 4 pilihan jawaban (A, B, C, D).
3. Berikan kunci jawaban beserta penjelasan singkat mengapa jawaban itu benar.
4. Output harus dalam format JSON murni.

Struktur JSON:
{{
  "judul_kuis": "Judul Kuis",
  "questions": [
    {{
      "no": 1,
      "pertanyaan": "Teks pertanyaan...",
      "options": {{
        "A": "Jawaban A",
        "B": "Jawaban B",
        "C": "Jawaban C",
        "D": "Jawaban D"
      }},
      "kunci_jawaban": "A",
      "penjelasan": "Karena..."
    }}
  ]
}}
"""

    try:
        # 2. Call AI
        print(f"DEBUG: Generating Quiz for {req.topik}...")
        response_text = await gemini_client.generate_content(prompt)
        
        # Clean JSON
        match = re.search(r'(\{.*\}|\[.*\])', response_text, re.DOTALL)
        if not match:
             raise HTTPException(status_code=500, detail="AI tidak memberikan format data yang benar.")
        
        quiz_data = json.loads(match.group(0))
        
        # 3. Save to DB
        new_quiz = SavedQuiz(
            user_id=user_id,
            mapel=req.mapel,
            topik=req.topik,
            tingkat_kesulitan=req.tingkat_kesulitan,
            quiz_data=quiz_data
        )
        db.add(new_quiz)
        await db.commit()
        await db.refresh(new_quiz)
        
        return {
            "status": "success",
            "quiz_id": new_quiz.id,
            "data": quiz_data
        }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Gagal generate soal: {str(e)}")

@router.post("/export-quiz-pdf")
async def export_quiz_pdf(req: ExportQuizRequest):
    try:
        pdf = FPDF()
        pdf.add_page()
        
        # Header
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, clean_text(f"Latihan Soal: {req.topik}"), ln=True, align='C')
        pdf.set_font("Arial", '', 12)
        pdf.cell(0, 10, clean_text(f"Mata Pelajaran: {req.mapel}"), ln=True, align='C')
        pdf.ln(5)
        
        # Questions
        questions = req.quiz_data.get("questions", [])
        for q in questions:
            pdf.set_font("Arial", 'B', 12)
            # Question text
            txt = clean_text(f"{q.get('no', '')}. {q.get('pertanyaan', '')}")
            pdf.multi_cell(0, 10, txt)
            
            pdf.set_font("Arial", '', 11)
            options = q.get("options", {})
            for key, val in options.items():
                pdf.cell(0, 8, clean_text(f"   {key}. {val}"), ln=True)
            pdf.ln(5)

        # Buffer
        pdf_bytes = pdf.output()
        safe_topik = re.sub(r'[^\w\s-]', '', req.topik).strip().replace(" ", "_")
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=Quiz_{safe_topik}.pdf",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Gagal export PDF Quiz: {str(e)}")

@router.post("/export-quiz-word")
async def export_quiz_word(req: ExportQuizRequest):
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading(f"Latihan Soal: {req.topik}", 0)
        subtitle = doc.add_paragraph(f"Mata Pelajaran: {req.mapel}")
        subtitle.alignment = 1 # Center
        
        questions = req.quiz_data.get("questions", [])
        for q in questions:
            p = doc.add_paragraph()
            p.add_run(f"{q.get('no', '')}. {q.get('pertanyaan', '')}").bold = True
            
            options = q.get("options", {})
            for key, val in options.items():
                doc.add_paragraph(f"   {key}. {val}")
        
        # Kunci Jawaban at the end
        doc.add_page_break()
        doc.add_heading("Kunci Jawaban & Penjelasan", level=1)
        for q in questions:
            p = doc.add_paragraph()
            p.add_run(f"No {q.get('no', '')}: ").bold = True
            p.add_run(f"{q.get('kunci_jawaban', '')}")
            
            expl = doc.add_paragraph()
            expl.add_run("Penjelasan: ").italic = True
            expl.add_run(f"{q.get('penjelasan', '')}")
            
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return Response(
            content=file_stream.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=Quiz_{req.topik}.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal export Word: {str(e)}")

# --- HISTORY ENDPOINTS ---

@router.post("/export-pdf")
async def export_rpp_pdf(req: ExportRPPRequest):
    try:
        # Robust input sanitization
        topik = str(req.topik or "Tanpa Judul")
        mapel = str(req.mapel or "Mata Pelajaran")
        kelas = str(req.kelas or "Semua")
        content_markdown = str(req.content_markdown or "")

        if topik.lower() == "null": topik = "Tanpa Judul"
        if mapel.lower() == "null": mapel = "Mata Pelajaran"
        if kelas.lower() == "null": kelas = "Semua"
        if content_markdown.lower() == "null": content_markdown = ""

        print(f"DEBUG: Exporting Synchronized PDF for {topik}...")
        pdf = FPDF()
        pdf.add_page()
        
        # --- HEADER (Matches Word/Modal) ---
        pdf.set_y(15)
        pdf.set_x(10)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Arial", 'B', 11)
        pdf.cell(190, 10, "MODUL AJAR (RPP)", ln=True, align='C')
        
        pdf.set_font("Arial", 'B', 18)
        pdf.set_x(10)
        pdf.multi_cell(190, 12, clean_text(clean_markdown_symbols(topik)), align='C')
        
        pdf.set_font("Arial", 'I', 10)
        pdf.set_x(10)
        meta_text_clean = clean_markdown_symbols(f"{mapel} | Kelas {kelas}")
        pdf.cell(190, 8, clean_text(meta_text_clean), ln=True, align='C')
        
        # Separator Line
        pdf.ln(2)
        pdf.set_draw_color(200, 200, 200)
        pdf.line(20, pdf.get_y(), 190, pdf.get_y())
        pdf.ln(10)
        
        # Content Parsing
        lines = content_markdown.split('\n')
        pdf.set_font("Arial", '', 11)
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Table detection
            is_table_start = '|' in line and i + 1 < len(lines) and re.match(r'^\s*\|?[:\-\s|]+\|?[:\-\s|]*\s*$', lines[i+1])
            if is_table_start:
                table_data = []
                header_line = line.strip().strip('|')
                headers = [clean_text(clean_markdown_symbols(c)) for c in header_line.split('|')]
                i += 2
                while i < len(lines):
                    row_line = lines[i].strip()
                    if not '|' in row_line and not row_line.startswith('|'): break
                    row_content = row_line.strip('|')
                    row = [clean_text(clean_markdown_symbols(c)) for c in row_content.split('|')]
                    if row:
                        while len(row) < len(headers): row.append("")
                        table_data.append(row[:len(headers)])
                    i += 1
                
                if headers or table_data:
                    pdf.ln(5)
                    pdf.set_line_width(0.2)
                    with pdf.table(width=190, padding=2, line_height=7) as table:
                        if headers:
                            header_row = table.row()
                            pdf.set_font("Arial", 'B', 10)
                            for h in headers: header_row.cell(h)
                        pdf.set_font("Arial", '', 10)
                        for r_data in table_data:
                            row = table.row()
                            for c in r_data: row.cell(c)
                    pdf.ln(5)
                continue

            if not line:
                pdf.ln(2)
                i += 1
                continue
            
            # Header handling
            if line.startswith('#'):
                clean_header = clean_markdown_symbols(re.sub(r'^#+\s*', '', line))
                if line.startswith('###'):
                    pdf.set_font("Arial", 'B', 12)
                    pdf.ln(3)
                elif line.startswith('##'):
                    pdf.set_font("Arial", 'B', 13)
                    pdf.ln(4)
                else: 
                    pdf.set_font("Arial", 'B', 14)
                    pdf.ln(5)
                
                pdf.set_x(10)
                pdf.multi_cell(0, 8, clean_text(clean_header))
                pdf.set_font("Arial", '', 11)
            
            # List handling (broader detection)
            elif re.match(r'^[\-\*•]\s*', line):
                # Extract text after any bullet-like symbol
                text = re.sub(r'^[\-\*•]\s*', '', line)
                text = clean_markdown_symbols(text)
                pdf.set_x(15)
                # Use a proper bullet character
                pdf.multi_cell(0, 7, clean_text(f"\x95 {text}")) # \x95 is bullet in many latin-1 encodings
                pdf.set_x(10)
            
            # Regular text
            else:
                text = clean_markdown_symbols(line)
                pdf.set_x(10)
                pdf.multi_cell(0, 7, clean_text(text))
            i += 1
        
        pdf_bytes = pdf.output()
        safe_topik = re.sub(r'[^\w\s-]', '', req.topik).strip().replace(" ", "_")
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=RPP_{safe_topik}.pdf",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Gagal export PDF RPP: {str(e)}")
@router.post("/export-word")
async def export_rpp_word(req: ExportRPPRequest):
    try:
        # Robust input sanitization
        topik = str(req.topik or "Tanpa Judul")
        mapel = str(req.mapel or "Mata Pelajaran")
        kelas = str(req.kelas or "Semua")
        content_markdown = str(req.content_markdown or "")

        if topik.lower() == "null": topik = "Tanpa Judul"
        if mapel.lower() == "null": mapel = "Mata Pelajaran"
        if kelas.lower() == "null": kelas = "Semua"
        if content_markdown.lower() == "null": content_markdown = ""

        print(f"DEBUG: Exporting Word RPP for {topik}...")
        doc = Document()
        
        # Title Section (Styled as Modal)
        h0 = doc.add_heading("MODUL AJAR (RPP)", 0)
        h0.alignment = 1
        for run in h0.runs: run.font.color.rgb = RGBColor(0, 0, 0)
        
        p_topik = doc.add_paragraph()
        p_topik.alignment = 1
        run_topik = p_topik.add_run(topik)
        run_topik.bold = True
        run_topik.font.size = Pt(18)
        run_topik.font.color.rgb = RGBColor(0, 0, 0)
        
        p_meta = doc.add_paragraph()
        p_meta.alignment = 1
        run_mapel = p_meta.add_run(f" {mapel} ")
        run_mapel.font.size = Pt(10)
        run_mapel.italic = True
        run_mapel.font.color.rgb = RGBColor(0, 0, 0)
        
        run_kelas = p_meta.add_run(f" | Kelas {kelas} ")
        run_kelas.font.size = Pt(10)
        run_kelas.italic = True
        run_kelas.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph("_" * 60).alignment = 1
        
        # Content Parsing with Table Support
        lines = content_markdown.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Table detection (More robust regex check)
            is_table_start = '|' in line and i + 1 < len(lines) and re.match(r'^\s*\|?[:\-\s|]+\|?[:\-\s|]*\s*$', lines[i+1])
            if is_table_start:
                header_line = line.strip().strip('|')
                headers = [c.strip().replace('**', '') for c in header_line.split('|')]
                
                i += 2 # Skip header and separator
                rows = []
                while i < len(lines):
                    row_line = lines[i].strip()
                    if not '|' in row_line and not row_line.startswith('|'):
                        break
                    
                    row_content = row_line.strip('|')
                    row = [c.strip().replace('**', '') for c in row_content.split('|')]
                    if row:
                        while len(row) < len(headers): row.append("")
                        rows.append(row[:len(headers)])
                    i += 1
                
                if headers or rows:
                    table = doc.add_table(rows=0, cols=len(headers))
                    table.style = 'Table Grid'
                    
                    if headers:
                        header_row = table.add_row().cells
                        for idx, hs in enumerate(headers):
                            p = header_row[idx].paragraphs[0]
                            r = p.add_run(hs)
                            r.bold = True
                            r.font.color.rgb = RGBColor(0, 0, 0)
                    
                    for row_data in rows:
                        cells = table.add_row().cells
                        for idx, val in enumerate(row_data):
                            p = cells[idx].paragraphs[0]
                            r = p.add_run(val)
                            r.font.color.rgb = RGBColor(0, 0, 0)
                continue

            if not line:
                i += 1
                continue
            
            # Header handling
            if line.startswith('#'):
                # Robustly remove any number of # at the start
                clean_header = re.sub(r'^#+\s*', '', line)
                level = 1
                if line.startswith('###'): level = 3
                elif line.startswith('##'): level = 2
                
                h = doc.add_heading(clean_header, level=level)
                # Set heading color to black
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
            
            # List handling
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                text = re.sub(r'^[\-\*]\s*', '', line)
                # Handle bold parts
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        r = p.add_run(part[2:-2])
                        r.bold = True
                    else:
                        r = p.add_run(part)
                    r.font.color.rgb = RGBColor(0, 0, 0)
            
            # Regular text
            else:
                p = doc.add_paragraph()
                # Handle bold parts
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        r = p.add_run(part[2:-2])
                        r.bold = True
                    else:
                        r = p.add_run(part)
                    r.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        safe_topik = re.sub(r'[^\w\s-]', '', req.topik).strip().replace(" ", "_")
        
        return Response(
            content=file_stream.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=RPP_{safe_topik}.docx",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except Exception as e:
        traceback.print_exc()

@router.get("/history")
async def get_rpp_history(
    user_id: int = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    from app.models.rpp_data import SavedRPP
    stmt = select(SavedRPP).where(SavedRPP.user_id == user_id).order_by(SavedRPP.created_at.desc())
    result = await db.execute(stmt)
    history = result.scalars().all()
    return history

@router.delete("/history/{rpp_id}")
async def delete_rpp(
    rpp_id: int,
    user_id: int = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    from app.models.rpp_data import SavedRPP
    stmt = select(SavedRPP).where(SavedRPP.id == rpp_id, SavedRPP.user_id == user_id)
    result = await db.execute(stmt)
    rpp = result.scalar_one_or_none()
    
    if not rpp:
        raise HTTPException(status_code=404, detail="RPP not found")
        
    await db.delete(rpp)
    await db.commit()
    return {"message": "RPP deleted successfully"}

@router.get("/quiz-history")
async def get_quiz_history(
    user_id: int = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    from app.models.rpp_data import SavedQuiz
    stmt = select(SavedQuiz).where(SavedQuiz.user_id == user_id).order_by(SavedQuiz.created_at.desc())
    result = await db.execute(stmt)
    history = result.scalars().all()
    return history

@router.delete("/quiz-history/{quiz_id}")
async def delete_quiz(
    quiz_id: int,
    user_id: int = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db)
):
    from app.models.rpp_data import SavedQuiz
    stmt = select(SavedQuiz).where(SavedQuiz.id == quiz_id, SavedQuiz.user_id == user_id)
    result = await db.execute(stmt)
    quiz = result.scalar_one_or_none()
    
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")
        
    await db.delete(quiz)
    await db.commit()
    return {"message": "Quiz deleted successfully"}
